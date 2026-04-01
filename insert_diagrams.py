"""
把生成的图表插入论文 docx
运行：python3 insert_diagrams.py
输出：于晨旭_毕业论文_高校外聘教师管理系统_含图版.docx
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIAGRAMS_DIR = '/Users/yuchenxu/Desktop/于晨旭毕设/diagrams'
SRC_DOCX = '/Users/yuchenxu/Desktop/于晨旭毕设/于晨旭_毕业论文_高校外聘教师管理系统.docx'
OUT_DOCX = '/Users/yuchenxu/Desktop/于晨旭毕设/于晨旭_毕业论文_高校外聘教师管理系统_含图版.docx'

# 图编号 → PNG 文件名（diagrams/ 目录下）
FIG_MAP = {
    '图3-1': ('fig3-1_roles.png',         12.0),
    '图4-1': ('fig4-1_architecture.png',  13.0),
    '图4-2': ('fig4-2_er.png',            15.0),
    '图4-3': ('fig4-3_statemachine.png',  13.0),
    '图5-1': ('fig5-1_jwt_sequence.png',  13.0),
}

doc = Document(SRC_DOCX)

# 找到所有包含占位标记的段落（先记录索引，再反向处理避免索引漂移）
# 遍历策略：找到【此处插入图X-X】段落 → 替换为图片 → 删除紧随其后的橙色提示行

replaced = 0
# 需要删除的 XML 元素（提示行）
to_remove = []

paras = doc.paragraphs
i = 0
while i < len(paras):
    para = paras[i]
    text = para.text.strip()

    matched_fig = None
    for fig_id, (fname, width_cm) in FIG_MAP.items():
        if f'【此处插入{fig_id}】' in text:
            matched_fig = (fig_id, fname, width_cm)
            break

    if matched_fig:
        fig_id, fname, width_cm = matched_fig
        img_path = os.path.join(DIAGRAMS_DIR, fname)

        if os.path.exists(img_path):
            # 清空占位段落，插入图片
            para.clear()
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, width=Cm(width_cm))
            replaced += 1
            print(f'  ✓ 插入 {fig_id}  ({fname})')

            # 下一段如果是橙色提示行，标记删除
            if i + 1 < len(paras):
                next_text = paras[i + 1].text
                if '[截图提示' in next_text or '[截图说明' in next_text or next_text.startswith('['):
                    to_remove.append(paras[i + 1]._element)
        else:
            print(f'  ✗ 图片文件不存在，跳过 {fig_id}: {img_path}')

    i += 1

# 删除提示行
for el in to_remove:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

doc.save(OUT_DOCX)
print(f'\n共插入 {replaced} 张图，已保存：')
print(f'  {OUT_DOCX}')
print()
print('剩余需手动插入的截图（系统运行截图）：')
remaining = [f for f in [
    '图3-2  系统用例图',
    '图5-2  教师信息管理页面',
    '图5-3  考勤管理页面',
    '图5-4  薪酬管理页面',
    '图5-5  审批时间线弹窗',
    '图5-6  合同管理页面',
    '图5-7  首页仪表盘',
    '图6-1  审批全流程测试',
]]
for r in remaining:
    print(f'  ○ {r}')
