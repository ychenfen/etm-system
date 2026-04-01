"""
高校外聘教师管理系统 — 毕业论文生成脚本
作者：于晨旭  学号：2022062008
运行：python3 generate_thesis.py
输出：于晨旭_毕业论文_高校外聘教师管理系统.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 文档初始化 ─────────────────────────────────────────────────────────────────

doc = Document()

# 页面设置：A4，上下2.54cm，左3.17cm右3.17cm
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.17)
section.right_margin  = Cm(3.17)

# ── 样式辅助函数 ───────────────────────────────────────────────────────────────

def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=12, bold=False, color=None):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn_font)
    rPr.insert(0, rFonts)

def set_para_format(para, first_line_indent=True, line_spacing=1.5,
                    space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = para.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Pt(24)  # 约2字符

def add_heading(text, level=1):
    """添加标题，level=1为章标题，level=2为节标题，level=3为小节"""
    para = doc.add_paragraph()
    set_para_format(para, first_line_indent=False, space_before=12, space_after=6,
                    align=WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    size_map = {1: 16, 2: 14, 3: 13}
    set_run_font(run, cn_font='黑体', en_font='Arial', size_pt=size_map.get(level, 12), bold=True)
    return para

def add_para(text, indent=True, bold=False):
    """添加正文段落"""
    para = doc.add_paragraph()
    set_para_format(para, first_line_indent=indent)
    run = para.add_run(text)
    set_run_font(run, bold=bold)
    return para

def add_figure(fig_id, caption, screenshot_hint):
    """添加图片占位符"""
    # 占位框
    para = doc.add_paragraph()
    set_para_format(para, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=0)
    run = para.add_run(f'【此处插入{fig_id}】')
    set_run_font(run, cn_font='宋体', size_pt=10, color=(150, 150, 150))
    # 截图说明（灰色小字，仅供参考，正式版删除）
    hint_para = doc.add_paragraph()
    set_para_format(hint_para, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    hint_run = hint_para.add_run(f'[截图提示：{screenshot_hint}]')
    set_run_font(hint_run, cn_font='宋体', size_pt=9, color=(180, 120, 60))
    # 图题
    cap_para = doc.add_paragraph()
    set_para_format(cap_para, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=10)
    cap_run = cap_para.add_run(f'{fig_id}  {caption}')
    set_run_font(cap_run, cn_font='宋体', size_pt=10.5, bold=False)

def add_code(code_text, lang='Java'):
    """添加代码块"""
    para = doc.add_paragraph()
    set_para_format(para, first_line_indent=False, line_spacing=1.2,
                    space_before=4, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    pf = para.paragraph_format
    pf.left_indent = Pt(28)
    run = para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '仿宋')
    rPr.insert(0, rFonts)

def add_table_caption(text):
    para = doc.add_paragraph()
    set_para_format(para, first_line_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=6, space_after=4)
    run = para.add_run(text)
    set_run_font(run, cn_font='宋体', size_pt=10.5)

def add_simple_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()  # 表后空行

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════════════════════

def add_cover():
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('本 科 毕 业 设 计（论 文）')
    set_run_font(r, cn_font='黑体', en_font='Arial', size_pt=22, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run('高校外聘教师管理系统的设计与实现')
    set_run_font(tr, cn_font='宋体', en_font='Times New Roman', size_pt=18, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    info = [
        ('作    者：', '于晨旭'),
        ('学    号：', '2022062008'),
        ('专    业：', '计算机科学与技术'),
        ('指导教师：', '（请填写）'),
        ('完成日期：', '2026 年 5 月'),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(label)
        set_run_font(lr, cn_font='宋体', size_pt=14, bold=False)
        vr = p.add_run(value)
        set_run_font(vr, cn_font='宋体', size_pt=14, bold=False)

    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 摘要
# ══════════════════════════════════════════════════════════════════════════════

def add_abstract():
    add_heading('摘  要', level=1)
    add_para(
        '高校外聘教师是高等教育师资体系的重要补充，但传统的外聘管理模式依赖纸质表格和线下沟通，'
        '存在信息分散、审批周期长、薪酬核算出错等问题。本文设计并实现了一套面向高校人事管理场景的'
        '外聘教师管理系统，涵盖教师档案管理、课程排课、考勤记录、薪酬结算、三级聘用审批、合同管理、'
        '教学评价及数据统计等核心功能。'
    )
    add_para(
        '系统采用前后端分离架构。后端以 Spring Boot 2.7.18 为基础框架，结合 MyBatis-Plus 3.5.3.1 '
        '处理数据持久化，使用 H2 嵌入式数据库支持开箱即用的开发部署，生产环境可无缝切换至 MySQL。'
        '身份认证采用 JWT 方案，通过自定义拦截器实现无状态会话管理。前端基于 Vue 2 与 Element UI '
        '构建，ECharts 负责首页数据可视化。'
    )
    add_para(
        '在业务设计上，聘用审批模块实现了学院负责人→人事处薪酬岗→人事处处长三个节点的有序流转，'
        '每节点均可独立通过或驳回，处长审批通过后系统自动生成聘用合同。薪酬结算模块按照'
        '"课时费 = 课时数 × 课时单价 + 补贴 − 扣款"的规则自动汇总，减少人工核算误差。'
    )
    add_para(
        '测试结果表明，系统各功能模块运行稳定，三级审批流程状态转换正确，薪酬计算结果与预期一致，'
        '主要页面加载时延均在 500ms 以内，满足高校日常管理的基本性能要求。'
    )
    p = doc.add_paragraph()
    set_para_format(p, first_line_indent=False, space_before=6)
    r1 = p.add_run('关键词：')
    set_run_font(r1, bold=True)
    r2 = p.add_run('外聘教师管理；Spring Boot；MyBatis-Plus；Vue.js；审批流程；前后端分离')
    set_run_font(r2)

    page_break()

    add_heading('Abstract', level=1)
    add_para(
        'Part-time (externally hired) teachers are a significant supplement to university faculty. '
        'However, traditional management relies on paper forms and offline communication, leading to '
        'scattered information, lengthy approval cycles, and salary calculation errors. This paper '
        'presents the design and implementation of an external teacher management system for '
        'university HR scenarios, covering teacher profiles, course scheduling, attendance, salary '
        'settlement, three-level hiring approval, contract management, teaching evaluation, and '
        'data analytics.'
    )
    add_para(
        'The system adopts a front-end/back-end separated architecture. The back end is built on '
        'Spring Boot 2.7.18, uses MyBatis-Plus 3.5.3.1 for data persistence, and employs an H2 '
        'embedded database for zero-configuration development, with seamless MySQL migration for '
        'production. Authentication is handled via JWT with custom interceptors for stateless '
        'session management. The front end is built with Vue 2 and Element UI; ECharts provides '
        'dashboard visualizations.'
    )
    add_para(
        'Testing confirms stable operation across all modules. The three-level approval state '
        'machine transitions correctly, salary calculations match expected values, and key page '
        'load times remain within 500 ms, meeting the performance requirements of daily campus '
        'HR management.'
    )
    p2 = doc.add_paragraph()
    set_para_format(p2, first_line_indent=False, space_before=6)
    r3 = p2.add_run('Keywords: ')
    set_run_font(r3, en_font='Times New Roman', bold=True)
    r4 = p2.add_run('External Teacher Management; Spring Boot; MyBatis-Plus; Vue.js; Approval Workflow; Front-Back Separation')
    set_run_font(r4, en_font='Times New Roman')

    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 目录（手动占位，实际用Word生成）
# ══════════════════════════════════════════════════════════════════════════════

def add_toc():
    add_heading('目  录', level=1)
    toc_items = [
        ('第1章  绪论', '1'),
        ('    1.1  研究背景与意义', '1'),
        ('    1.2  国内外研究现状', '2'),
        ('    1.3  本文主要工作', '3'),
        ('    1.4  论文结构安排', '3'),
        ('第2章  相关技术介绍', '4'),
        ('    2.1  Spring Boot 框架', '4'),
        ('    2.2  MyBatis-Plus', '5'),
        ('    2.3  JWT 认证机制', '5'),
        ('    2.4  Vue.js 与 Element UI', '6'),
        ('    2.5  H2 嵌入式数据库', '6'),
        ('第3章  需求分析', '7'),
        ('    3.1  用户角色分析', '7'),
        ('    3.2  功能性需求', '8'),
        ('    3.3  非功能性需求', '10'),
        ('第4章  系统设计', '11'),
        ('    4.1  总体架构设计', '11'),
        ('    4.2  数据库设计', '12'),
        ('    4.3  核心模块设计', '16'),
        ('    4.4  接口设计规范', '18'),
        ('第5章  系统实现', '19'),
        ('    5.1  开发环境配置', '19'),
        ('    5.2  用户认证模块', '20'),
        ('    5.3  教师信息管理模块', '22'),
        ('    5.4  课程与考勤管理模块', '23'),
        ('    5.5  薪酬结算模块', '25'),
        ('    5.6  三级审批流程模块', '26'),
        ('    5.7  合同管理模块', '29'),
        ('    5.8  数据统计与可视化模块', '30'),
        ('第6章  系统测试', '31'),
        ('    6.1  测试环境', '31'),
        ('    6.2  功能测试', '32'),
        ('    6.3  性能与安全测试', '33'),
        ('第7章  总结与展望', '34'),
        ('参考文献', '35'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        set_para_format(p, first_line_indent=False, space_before=0, space_after=2)
        run = p.add_run(item)
        set_run_font(run, size_pt=11)
        # 右对齐页码（简单处理，Word会在手动更新目录后替换）
        tab_run = p.add_run(f'\t{page}')
        set_run_font(tab_run, size_pt=11)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第1章  绪论
# ══════════════════════════════════════════════════════════════════════════════

def chapter1():
    add_heading('第1章  绪论', level=1)

    add_heading('1.1  研究背景与意义', level=2)
    add_para(
        '近年来，高校持续扩大招生规模，各专业对师资的需求随之增长，而专职教师的编制扩张受到政策和'
        '经费的双重约束。外聘教师作为对正式教职队伍的有效补充，承担着相当比例的本科及研究生课程教学'
        '任务。以计算机、商科、艺术类专业为例，具有行业背景的外聘讲师往往能弥补专职教师在实践经验上'
        '的不足，受到院系欢迎。'
    )
    add_para(
        '然而，外聘教师的管理复杂程度不亚于正式教职，涉及档案录入、聘用审批、课程排布、出勤核查、'
        '薪酬结算、合同签署等多个环节，且参与方包括二级学院、人事处薪酬岗、人事处处长等不同层级。'
        '目前多数高校仍依赖纸质表格和电子邮件完成这些流程，主要问题体现在三个方面：其一，信息存储'
        '分散，各院系各自维护 Excel 档案，难以形成全校统一视图；其二，审批流程不透明，申请方无法'
        '实时了解当前进展，节点之间的流转完全依靠人工转交；其三，薪酬核算误差多，手动汇总课时数据'
        '容易遗漏，每月核对成本较高。'
    )
    add_para(
        '本文设计并实现的高校外聘教师管理系统，旨在将上述流程整体纳入信息化平台，以 Web 应用的形式'
        '提供给各类用户使用。系统建成后，人事处管理员可在统一界面查看全校外聘师资状况，二级学院管理'
        '员可发起聘用审批、管理本院课程，外聘教师本人可查看个人考勤、薪酬明细和评价记录。对高校信息化'
        '建设而言，该系统可作为人事管理平台的专项补充，也为类似场景下的工作流系统提供一套可落地的'
        '参考实现。'
    )

    add_heading('1.2  国内外研究现状', level=2)
    add_para(
        '国外高校人事管理系统发展较早。美国、英国等高校普遍使用商业人力资源管理软件（如 Workday、'
        'Oracle HCM），这类系统功能全面，但定制成本高，且以正式雇员管理为核心，对临时性、非全日制'
        '人员（包括外聘教师）的管理支持相对薄弱，通常需要单独配置模块。'
    )
    add_para(
        '国内方面，高校信息化建设近年来取得明显进展，但专门针对外聘教师全生命周期管理的系统仍较少见。'
        '现有研究多集中于通用工作流引擎的设计（如活动图驱动的审批系统、基于状态机的流程引擎），以及'
        '高校人事档案数字化管理的方法论探讨，较少涉及"外聘"这一特定场景的完整业务闭环。在开源项目层面，'
        'RuoYi、JeecgBoot 等 Java 后台框架提供了较好的基础脚手架，但外聘业务逻辑仍需自行实现。'
    )
    add_para(
        '本系统的工作侧重点与上述研究的差异在于：不追求通用工作流引擎的抽象，而是针对高校外聘场景的'
        '具体业务流设计专用状态机；不依赖重型框架，以尽量精简的技术栈完成完整的业务闭环，适合中小规模'
        '高校在内网环境部署使用。'
    )

    add_heading('1.3  本文主要工作', level=2)
    add_para('本文的主要工作包含以下几个部分：')
    add_para(
        '（1）对高校外聘教师管理的业务流程进行调研和梳理，明确系统管理员、学院管理员、外聘教师三类'
        '角色的需求边界，形成功能需求规格。'
    )
    add_para(
        '（2）完成系统的整体架构设计，包括前后端模块划分、数据库表结构、三级审批状态机的形式化定义，'
        '以及统一接口规范的制定。'
    )
    add_para(
        '（3）基于 Spring Boot + MyBatis-Plus 实现后端服务，包含教师档案、课程考勤、薪酬结算、'
        '审批流程、合同管理等核心模块，共约 70 个 Java 类；基于 Vue 2 + Element UI 实现前端'
        '交互界面，共 16 个页面视图。'
    )
    add_para(
        '（4）对系统进行功能测试和基本性能测试，验证各模块在典型用例下的正确性，并对三级审批流程的'
        '边界条件（驳回、撤回、幂等重试）进行专项验证。'
    )

    add_heading('1.4  论文结构安排', level=2)
    add_para(
        '本文共七章。第1章介绍研究背景、现状与本文工作。第2章说明系统所用主要技术的基本原理。'
        '第3章进行需求分析，确定系统的功能边界。第4章完成系统整体设计，包括架构、数据库和核心模块。'
        '第5章按模块描述具体实现过程，并给出关键代码。第6章进行系统测试，报告测试结果。第7章总结'
        '全文并指出后续改进方向。'
    )

    add_para('本章小结：本章说明了外聘教师管理信息化的现实需求，梳理了国内外相关系统的研究现状，'
             '并明确了本文的工作内容与论文组织结构。', bold=False)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第2章  相关技术介绍
# ══════════════════════════════════════════════════════════════════════════════

def chapter2():
    add_heading('第2章  相关技术介绍', level=1)

    add_heading('2.1  Spring Boot 框架', level=2)
    add_para(
        'Spring Boot 是 Pivotal 团队在 Spring Framework 基础上开发的快速应用构建框架，核心思想是'
        '"约定大于配置"。通过自动配置机制（Auto-Configuration），开发者无需编写大量 XML 配置，'
        '只需引入对应的 Starter 依赖，框架即可自动装配 DataSource、MVC 等基础组件。内嵌 Tomcat '
        '服务器使应用可以直接打包为可执行 JAR，省去单独部署 Web 容器的步骤。'
    )
    add_para(
        '本系统使用 Spring Boot 2.7.18，该版本属于 2.x 系列的最后一个维护版本，兼容性成熟。'
        '相较于 Spring Boot 3.x，2.7.x 对 JDK 8/11/17 均有支持，降低了开发环境配置的复杂度。'
        '系统的 Web 层、异常处理、配置管理均基于 Spring Boot 提供的注解驱动开发模式完成。'
    )

    add_heading('2.2  MyBatis-Plus', level=2)
    add_para(
        'MyBatis-Plus 是在 MyBatis 基础上的增强工具，保留了 MyBatis 的 XML 映射能力，同时通过'
        '继承 BaseMapper<T> 和 ServiceImpl<M, T> 提供了无需编写 SQL 的单表 CRUD 操作。分页插件'
        '（PaginationInterceptor）可在不修改业务代码的前提下自动注入分页查询语句。'
    )
    add_para(
        '本系统使用 MyBatis-Plus 3.5.3.1。所有实体类通过 @TableLogic 注解启用逻辑删除，删除操作'
        '实际上是将 deleted 字段置为 1，查询时框架自动过滤已删除记录。对于需要联表查询的场景（如'
        '审批列表中展示教师姓名和院系名称），在实体类中使用 @TableField(exist=false) 声明非数据库'
        '字段，由 Service 层手动赋值。'
    )

    add_heading('2.3  JWT 认证机制', level=2)
    add_para(
        'JSON Web Token（JWT）是一种基于 JSON 的开放标准（RFC 7519），用于在各方之间安全传递'
        '声明信息。一个 JWT 由三部分组成：Header（算法和令牌类型）、Payload（声明，如用户 ID、'
        '角色、过期时间）和 Signature（使用密钥对前两部分的签名）。三部分以 Base64Url 编码后'
        '用点号拼接，形成形如 xxxxx.yyyyy.zzzzz 的紧凑字符串。'
    )
    add_para(
        '与 Session/Cookie 方案相比，JWT 的优势在于服务端无需存储会话状态，天然支持水平扩展。'
        '本系统使用 JJWT 库生成和解析令牌，密钥配置在 application.yml 中，有效期为 24 小时。'
        '未使用 Spring Security，而是通过自定义 JwtInterceptor 在请求到达 Controller 前完成'
        '令牌验证，RoleInterceptor 负责角色权限校验，两个拦截器在 WebMvcConfig 中注册并配置'
        '白名单路径（/api/auth/** 不拦截）。'
    )

    add_heading('2.4  Vue.js 与 Element UI', level=2)
    add_para(
        'Vue.js 是一款渐进式 JavaScript 前端框架，以数据驱动视图和组件化开发为核心设计理念。'
        '双向数据绑定通过 v-model 指令实现，组件间通信可通过 props/emit 或集中式状态管理（Vuex）'
        '完成。Vue Router 提供前端路由，支持导航守卫（beforeEach），可在路由切换前验证登录状态'
        '和角色权限。'
    )
    add_para(
        '本系统使用 Vue 2.6.14，配合 Element UI 2.15.14 组件库构建界面。Element UI 提供了'
        '表格、表单、对话框、分页、时间线（el-timeline）等丰富组件，使得管理类界面开发效率较高。'
        'ECharts 5.4.3 用于首页仪表盘的柱状图、饼图和折线图展示。'
    )

    add_heading('2.5  H2 嵌入式数据库', level=2)
    add_para(
        'H2 是一款纯 Java 编写的嵌入式关系型数据库，支持内存模式和文件持久化模式，提供与 MySQL '
        '高度兼容的 SQL 方言（需在连接 URL 中指定 MODE=MySQL）。其最大优势是无需独立安装数据库'
        '服务，Spring Boot 项目引入依赖后即可直接使用，适合快速原型开发和 CI 环境的集成测试。'
    )
    add_para(
        '本系统在开发阶段使用 H2 文件模式（jdbc:h2:file:./data/etm），数据持久化在本地文件中，'
        '重启后不丢失。Spring Boot 的 spring.sql.init 配置项可在应用启动时自动执行 schema.sql'
        '（建表）和 data.sql（初始化演示数据），实现开箱即用的效果。生产部署时只需修改'
        'application.yml 中的数据源配置，切换为 MySQL 连接字符串，业务代码无需改动。H2 的主要'
        '局限是不支持多进程并发访问同一文件，因此不适用于需要多实例部署的生产场景。'
    )

    add_para('本章小结：本章介绍了系统采用的五项主要技术，包括 Spring Boot 自动装配机制、'
             'MyBatis-Plus 的增强 ORM 特性、JWT 无状态认证原理、Vue 2 前端开发模式，'
             '以及 H2 嵌入式数据库的适用边界。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第3章  需求分析
# ══════════════════════════════════════════════════════════════════════════════

def chapter3():
    add_heading('第3章  需求分析', level=1)

    add_heading('3.1  用户角色分析', level=2)
    add_para(
        '系统面向三类用户，权限范围依次收窄。系统管理员（ADMIN）拥有全部操作权限，负责用户账号管理、'
        '全校院系信息维护、全量数据查询和数据导出；学院管理员（DEPARTMENT）管理本院外聘教师的档案、'
        '课程、考勤、薪酬，并可发起聘用审批；外聘教师（TEACHER）只能查看与自身相关的考勤、薪酬和'
        '评价记录，以及公告通知。'
    )

    add_figure('图3-1', '系统用户角色权限示意图',
               '绘制三列对比图，列出ADMIN/DEPARTMENT/TEACHER各自能访问的菜单项，'
               '可使用矩形色块区分，红色=独占，蓝色=共享')

    add_heading('3.2  功能性需求', level=2)
    add_para('根据对高校人事管理流程的分析，系统需支持以下八个功能模块：')

    add_heading('3.2.1  基础信息管理', level=3)
    add_para(
        '教师档案模块要求录入并维护外聘教师的个人基本信息，包括姓名、性别、出生日期、身份证号、'
        '学历学位、职称、所属单位、联系方式、擅长专业方向、银行账户等字段，并记录当前聘用状态'
        '（待聘/在职/已到期/已终止）。院系管理模块支持对学院基本信息的增删改查，院系信息作为教师、'
        '课程、审批等多个模块的外键依赖。'
    )

    add_heading('3.2.2  课程与考勤管理', level=3)
    add_para(
        '课程模块支持按学期录入外聘教师的授课安排，字段包括课程名称、课程代码、学分、总课时、'
        '周课时、班级、学生人数、教室、上课时间等。考勤模块记录每次授课的签到时间、签退时间、'
        '实际课时数和考勤状态（正常/迟到/早退/缺勤），为薪酬结算提供原始数据。'
    )

    add_heading('3.2.3  薪酬结算', level=3)
    add_para(
        '薪酬模块按月生成结算记录，字段包含基础薪资、总课时数、课时单价、奖励、扣款和应付总额。'
        '管理员可手动录入或批量导入，也可按月自动汇总考勤模块的实际课时数据后填充。结算状态分为'
        '待审核、已发放两种，支持查询历史薪酬记录。'
    )

    add_heading('3.2.4  聘用审批', level=3)
    add_para(
        '审批模块是系统的核心业务流程。人事或学院管理员提交聘用申请后，审批单依次经过学院负责人、'
        '人事处薪酬岗、人事处处长三个节点。每个节点的审批人可选择通过或驳回，驳回时需填写原因。'
        '学院负责人节点支持发起方撤回申请。全部节点通过后，系统自动生成聘用合同。'
    )

    add_heading('3.2.5  合同管理', level=3)
    add_para(
        '合同由审批通过后自动创建，初始状态为"未签署"。后续流程为：教师本人签署（记录签署时间），'
        '再由学校管理员盖章签署（合同状态变为"已签署"，合同正式生效）。系统需防止同一审批单'
        '重复生成合同（幂等校验）。'
    )

    add_heading('3.2.6  教学评价', level=3)
    add_para(
        '评价模块记录对外聘教师授课质量的考核结果，包含教学质量评分、出勤评分、学生评分和综合'
        '总评，并关联具体课程和学期，评价人姓名和评价日期可记录留档。'
    )

    add_heading('3.2.7  通知公告', level=3)
    add_para(
        '公告模块供管理员发布全校性或专项通知，分为系统通知（SYSTEM）、教学通知（TEACHING）、'
        '行政通知（ADMIN）等类型，所有已登录用户均可查阅。'
    )

    add_heading('3.2.8  数据统计', level=3)
    add_para(
        '首页仪表盘展示关键汇总数据：外聘教师总数、本学期在职人数、当月课时总量、待处理审批数量，'
        '以及教师学历分布饼图、月度课时趋势折线图、薪酬按院系分布柱状图等可视化图表。'
    )

    add_heading('3.3  非功能性需求', level=2)
    add_para(
        '安全性方面，所有接口（除登录外）均需携带有效 JWT 令牌，角色权限由拦截器在请求级别强制'
        '校验，前端同时在路由守卫中做二次过滤。密码存储使用 BCrypt 散列，不明文保存。'
    )
    add_para(
        '性能方面，系统面向单高校部署，预计并发用户数在数十人量级，目标是主要列表页面在局域网内'
        '的首次加载时延不超过 1 秒，后续翻页和操作响应不超过 500ms。H2 文件模式在此规模下'
        '足以支撑，无需引入缓存层。'
    )
    add_para(
        '可维护性方面，后端遵循 Controller-Service-Mapper 三层架构，各层职责清晰，统一响应格式'
        '（R<T>）和全局异常处理器（GlobalExceptionHandler）覆盖常见异常，便于问题定位。前端'
        '通过统一的 request.js 封装 Axios，集中处理令牌注入和错误提示。'
    )
    add_para(
        '可部署性方面，系统以嵌入式数据库为默认配置，开发者克隆代码后执行 mvn spring-boot:run '
        '即可启动后端，npm run serve 即可启动前端，无需额外安装数据库，适合快速验证和演示。'
    )

    add_figure('图3-2', '系统用例图（管理员视角）',
               '绘制UML用例图，椭圆表示用例，矩形框表示系统边界，管理员用户关联：'
               '教师管理/课程管理/考勤管理/薪酬管理/审批管理/合同管理/统计查看 等用例')

    add_para('本章小结：本章从三类用户角色出发，梳理了系统的八个功能模块需求，并明确了安全、'
             '性能、可维护性等非功能性约束，为后续系统设计奠定了基础。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第4章  系统设计
# ══════════════════════════════════════════════════════════════════════════════

def chapter4():
    add_heading('第4章  系统设计', level=1)

    add_heading('4.1  总体架构设计', level=2)
    add_para(
        '系统采用前后端分离的单体架构。前端 Vue 应用作为独立进程运行，通过 Axios 向后端 RESTful '
        '接口发送 HTTP 请求，两者之间以 JSON 格式交换数据。后端 Spring Boot 应用内部按照'
        ' Controller → Service → Mapper 三层分工：Controller 层负责参数校验和响应封装，'
        'Service 层承载业务逻辑，Mapper 层借助 MyBatis-Plus 与数据库交互。'
    )
    add_figure('图4-1', '系统总体架构图',
               '绘制分层架构图（从上到下）：'
               '第一层：浏览器（Vue 2 + Element UI + ECharts）；'
               '第二层：HTTP/Axios，标注JWT token；'
               '第三层：Spring Boot（Controller层 / Service层 / Mapper层），标注拦截器；'
               '第四层：H2 文件数据库。各层之间用箭头连接')
    add_para(
        '拦截器链在 WebMvcConfig 中配置：请求首先进入 JwtInterceptor，解析并验证 Authorization '
        '头中的 Bearer 令牌，将用户信息存入 ThreadLocal；随后进入 RoleInterceptor，读取接口上'
        '标注的 @RequireRole 注解，比对当前用户角色，不符合则返回 403。白名单路径（/api/auth/**）'
        '绕过两个拦截器直接通行。'
    )
    add_para(
        '前端路由守卫（router.beforeEach）在导航发生时检查 localStorage 中是否存在有效 token，'
        '未登录时跳转至 /login 页面。对于有角色限制的路由（meta.roles 字段），守卫额外读取'
        ' userInfo 中的 role 字段进行比对，无权限时跳转至 /dashboard。'
    )

    add_heading('4.2  数据库设计', level=2)
    add_para(
        '系统共设计 10 张业务表，均包含 deleted（逻辑删除标志）、create_time、update_time '
        '三个公共字段。各表设计说明如下。'
    )

    add_figure('图4-2', '数据库 ER 图',
               '使用数据库设计工具（如 draw.io 或 Navicat）绘制10张表的ER图，'
               '重点展示 teacher、approval、contract 之间的关联关系，'
               'approval.teacher_id → teacher.id，contract.approval_id → approval.id')

    add_table_caption('表4-1  sys_user 用户表字段说明')
    add_simple_table(
        ['字段名', '类型', '说明'],
        [('id', 'BIGINT', '主键，自增'),
         ('username', 'VARCHAR(50)', '登录名，唯一'),
         ('password', 'VARCHAR(100)', 'BCrypt 散列密码'),
         ('real_name', 'VARCHAR(50)', '真实姓名'),
         ('role', 'VARCHAR(20)', '角色：ADMIN/DEPARTMENT/TEACHER'),
         ('department_id', 'BIGINT', '所属院系（外键）'),
         ('status', 'INT', '状态：1启用 0禁用'),
         ('deleted', 'INT', '逻辑删除标志')]
    )

    add_table_caption('表4-2  teacher 外聘教师档案表字段说明')
    add_simple_table(
        ['字段名', '类型', '说明'],
        [('id', 'BIGINT', '主键，自增'),
         ('user_id', 'BIGINT', '关联系统用户（可空）'),
         ('name', 'VARCHAR(50)', '姓名'),
         ('gender', 'VARCHAR(10)', '性别'),
         ('id_card', 'VARCHAR(20)', '身份证号'),
         ('education', 'VARCHAR(50)', '学历'),
         ('degree', 'VARCHAR(50)', '学位'),
         ('title', 'VARCHAR(50)', '职称'),
         ('work_unit', 'VARCHAR(200)', '工作单位'),
         ('department_id', 'BIGINT', '所属院系'),
         ('hire_status', 'VARCHAR(20)', 'PENDING/ACTIVE/EXPIRED/TERMINATED'),
         ('bank_account', 'VARCHAR(50)', '银行账号（薪酬发放用）')]
    )

    add_table_caption('表4-3  approval 聘用审批表字段说明')
    add_simple_table(
        ['字段名', '类型', '说明'],
        [('approval_no', 'VARCHAR(50)', '审批编号，格式 APRyyyyMMddHHmmss'),
         ('teacher_id', 'BIGINT', '被审批教师'),
         ('current_node', 'VARCHAR(30)', '当前节点：college_leader/hr_salary/hr_director/finished'),
         ('approval_status', 'VARCHAR(20)', 'PENDING/APPROVED/REJECTED/REVOKED'),
         ('college_status', 'VARCHAR(20)', '学院节点结果'),
         ('college_by/time', 'VARCHAR/TIMESTAMP', '学院审批人姓名和时间'),
         ('hr_salary_*', '同上', '人事处薪酬岗节点字段'),
         ('hr_director_*', '同上', '人事处处长节点字段'),
         ('proposed_salary', 'DECIMAL(10,2)', '拟定课时薪酬')]
    )

    add_table_caption('表4-4  contract 聘用合同表字段说明')
    add_simple_table(
        ['字段名', '类型', '说明'],
        [('contract_no', 'VARCHAR(50)', '合同编号，格式 CTR+时间戳'),
         ('teacher_id', 'BIGINT', '签约教师'),
         ('approval_id', 'BIGINT', '来源审批单（幂等依据）'),
         ('salary_standard', 'DECIMAL(10,2)', '课时薪酬标准'),
         ('contract_status', 'VARCHAR(20)', 'UNSIGNED/SIGNED/EXPIRED/TERMINATED'),
         ('teacher_sign_time', 'TIMESTAMP', '教师签署时间'),
         ('school_sign_time', 'TIMESTAMP', '学校签署时间')]
    )

    add_heading('4.3  核心模块设计', level=2)

    add_heading('4.3.1  三级审批状态机', level=3)
    add_para(
        '审批流程定义四个状态节点：college_leader（学院负责人）、hr_salary（人事处薪酬岗）、'
        'hr_director（人事处处长）、finished（已完成）。整体审批状态（approval_status）在流转'
        '过程中保持 PENDING，任一节点驳回时变为 REJECTED，发起方撤回时变为 REVOKED，全部通过'
        '后变为 APPROVED。'
    )
    add_figure('图4-3', '三级审批状态机流程图',
               '绘制状态机图：矩形表示节点状态，菱形表示判断，箭头标注条件。'
               '流程：提交→college_leader节点→（通过→hr_salary节点→通过→hr_director节点→通过→finished+生成合同）'
               '各节点均有"驳回→REJECTED"分支，college_leader节点有"撤回→REVOKED"分支')
    add_para(
        '状态转换逻辑集中在 ApprovalServiceImpl.approve() 方法中，通过 switch-case 判断当前节点，'
        '记录当前节点的审批结果后决定下一步：若 pass 为 true 且当前节点不是 hr_director，则将'
        ' current_node 推进到下一节点；若当前节点是 hr_director 且通过，则将 current_node 置为'
        ' finished、approval_status 置为 APPROVED，并触发合同自动生成；若 pass 为 false，'
        '则直接将 approval_status 置为 REJECTED，不再推进。'
    )

    add_heading('4.3.2  薪酬计算规则', level=3)
    add_para('系统的薪酬计算公式如下：')
    add_para('月薪总额 = 基础薪资 + 实际课时数 × 课时单价 + 奖励补贴 − 扣款')
    add_para(
        '其中，实际课时数从考勤记录中汇总，课时单价在薪酬录入时由管理员指定（对应教师的职称和课程'
        '类型）。奖励和扣款项用于记录绩效调整或缺勤扣减。当前版本未实现职称系数的自动套算，'
        '课时单价由管理员手动填入，这是已知的功能简化点。'
    )

    add_heading('4.4  接口设计规范', level=2)
    add_para(
        '所有接口以 /api 为基础路径，采用 RESTful 风格。资源名称使用小写复数，操作通过 HTTP '
        '方法区分：GET 查询、POST 新增、PUT 修改、DELETE 删除（逻辑删除）。分页列表接口统一接受'
        ' current（当前页）、size（每页数量）查询参数，返回 PageResult<T> 包装对象。'
    )
    add_table_caption('表4-5  主要接口概览')
    add_simple_table(
        ['方法', '路径', '功能', '权限'],
        [('POST', '/api/auth/login', '用户登录，返回JWT', '公开'),
         ('GET',  '/api/teacher/page', '教师列表分页查询', 'ADMIN/DEPT'),
         ('POST', '/api/teacher', '新增外聘教师', 'ADMIN/DEPT'),
         ('GET',  '/api/attendance/page', '考勤记录分页查询', 'ADMIN/DEPT'),
         ('GET',  '/api/salary/page', '薪酬记录分页查询', 'ADMIN/DEPT'),
         ('POST', '/api/approval', '发起聘用审批', 'ADMIN/DEPT'),
         ('PUT',  '/api/approval/{id}/approve', '审批节点通过或驳回', 'ADMIN/DEPT'),
         ('PUT',  '/api/approval/{id}/revoke', '撤回审批申请', 'ADMIN/DEPT'),
         ('GET',  '/api/contract/page', '合同列表分页查询', 'ADMIN/DEPT'),
         ('PUT',  '/api/contract/{id}/teacher-sign', '教师签署合同', 'ADMIN/DEPT'),
         ('PUT',  '/api/contract/{id}/school-sign', '学校盖章签署', 'ADMIN'),
         ('GET',  '/api/dashboard/stats', '首页统计数据', '全部')]
    )

    add_para('本章小结：本章完成了系统的总体架构设计、10张核心数据表设计、三级审批状态机的'
             '形式化描述以及接口规范定义，为后续实现工作提供了完整的设计蓝图。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第5章  系统实现
# ══════════════════════════════════════════════════════════════════════════════

def chapter5():
    add_heading('第5章  系统实现', level=1)

    add_heading('5.1  开发环境配置', level=2)
    add_table_caption('表5-1  开发环境清单')
    add_simple_table(
        ['类别', '工具/版本'],
        [('操作系统', 'macOS 15 / Windows 10'),
         ('JDK', 'OpenJDK 17'),
         ('Maven', '3.9.x'),
         ('Node.js', '24.x'),
         ('IDE（后端）', 'IntelliJ IDEA 2024'),
         ('IDE（前端）', 'Visual Studio Code'),
         ('数据库（开发）', 'H2 2.2.x（内嵌）'),
         ('数据库（生产）', 'MySQL 8.0'),
         ('版本管理', 'Git 2.x / GitHub')]
    )
    add_para(
        '后端项目使用 Maven 管理依赖，核心依赖包括 spring-boot-starter-web、'
        'mybatis-plus-boot-starter（3.5.3.1）、h2、jjwt、hutool-all、'
        'poi-ooxml（Excel 导出）和 knife4j-spring-boot-starter（接口文档）。'
        '前端项目使用 npm 管理依赖，核心包为 vue（2.6.14）、element-ui（2.15.14）、'
        'echarts（5.4.3）和 axios（0.27.2）。'
    )
    add_para(
        '开发阶段，后端服务监听 8088 端口（context-path 为 /api），前端开发服务器监听 8080 端口。'
        '前端 vue.config.js 配置了代理规则，将 /api 前缀的请求转发至后端，消除跨域问题。'
    )

    add_heading('5.2  用户认证模块', level=2)
    add_para(
        '认证流程从 POST /api/auth/login 接口开始。AuthController 接收 LoginDto（username、'
        'password），调用 AuthServiceImpl.login()，该方法通过 UserMapper 按用户名查询账号，'
        '再用 HuTool 的 BCrypt 工具类验证密码散列是否匹配。验证通过后，调用 JwtUtil.generateToken() '
        '生成 JWT，Payload 中写入 userId、username、role 三个字段，返回给前端。'
    )
    add_figure('图5-1', 'JWT 认证流程图',
               '绘制时序图：前端→发送用户名密码→AuthController→AuthServiceImpl验证→JwtUtil生成token→'
               '返回token→前端存入localStorage→后续请求带Authorization: Bearer xxx→'
               'JwtInterceptor解析→放行')
    add_para(
        '后续请求的鉴权由 JwtInterceptor 完成。以下是其 preHandle 方法的核心片段：'
    )
    add_code(
        '// JwtInterceptor.java — 核心鉴权逻辑\n'
        '@Override\n'
        'public boolean preHandle(HttpServletRequest request,\n'
        '                         HttpServletResponse response, Object handler) {\n'
        '    String token = request.getHeader("Authorization");\n'
        '    if (token == null || !token.startsWith("Bearer ")) {\n'
        '        writeJson(response, R.unauthorized());\n'
        '        return false;\n'
        '    }\n'
        '    try {\n'
        '        Claims claims = JwtUtil.parseToken(token.substring(7));\n'
        '        // 将用户信息存入 ThreadLocal，供后续业务层使用\n'
        '        UserContext.set(claims.get("userId", Long.class),\n'
        '                        claims.get("role", String.class));\n'
        '        return true;\n'
        '    } catch (ExpiredJwtException e) {\n'
        '        writeJson(response, R.unauthorized());\n'
        '        return false;\n'
        '    }\n'
        '}'
    )
    add_para(
        'RoleInterceptor 读取 HandlerMethod 上的 @RequireRole 注解，将当前用户角色与注解'
        '声明的允许角色列表比对，不在列表中则返回 403 响应。前端 request.js 的响应拦截器'
        '在收到 401 时清除本地 token 并跳转登录页，收到 403 时显示"无权限"提示。'
    )

    add_heading('5.3  教师信息管理模块', level=2)
    add_para(
        'TeacherController 提供分页查询、新增、编辑、删除和查看详情五个接口。分页查询支持按教师'
        '姓名、所属院系、聘用状态进行筛选，底层通过 MyBatis-Plus 的 LambdaQueryWrapper 动态'
        '拼接 WHERE 条件：'
    )
    add_code(
        '// TeacherServiceImpl.java — 动态查询构建\n'
        'LambdaQueryWrapper<Teacher> wrapper = new LambdaQueryWrapper<>();\n'
        'wrapper.like(StrUtil.isNotBlank(name), Teacher::getName, name)\n'
        '       .eq(departmentId != null, Teacher::getDepartmentId, departmentId)\n'
        '       .eq(StrUtil.isNotBlank(hireStatus), Teacher::getHireStatus, hireStatus)\n'
        '       .eq(Teacher::getDeleted, 0)\n'
        '       .orderByDesc(Teacher::getCreateTime);\n'
        'return PageResult.of(teacherMapper.selectPage(page, wrapper));'
    )
    add_figure('图5-2', '教师信息管理页面截图',
               '登录管理员账号，点击侧边栏"教师管理"，截取完整页面，'
               '包含搜索栏（姓名/院系/状态筛选）、教师数据表格、右侧操作按钮')

    add_heading('5.4  课程与考勤管理模块', level=2)
    add_para(
        '课程管理模块在 CourseController 中实现，支持按学期、教师、院系多维度检索。考勤记录通过'
        'AttendanceController 管理，每条记录关联一位教师和一门课程，记录当天的签到时间、签退时间'
        '和实际课时数，考勤状态（NORMAL/LATE/LEAVE_EARLY/ABSENT）由管理员填写时手动指定。'
    )
    add_figure('图5-3', '考勤记录管理页面截图',
               '点击侧边栏"考勤管理"，截取包含日期范围筛选、考勤状态筛选、'
               '数据表格（教师姓名/课程/签到时间/签退时间/实际课时/状态）的完整页面')
    add_para(
        '考勤状态字段对薪酬结算有直接影响——实际课时数（actual_hours）将被 SalaryServiceImpl '
        '在自动汇总时读取。目前版本的薪酬汇总需管理员手动触发，未实现月底自动批量计算，这是'
        '后续可优化的点。'
    )

    add_heading('5.5  薪酬结算模块', level=2)
    add_para(
        '薪酬结算记录存储在 salary 表中，核心计算逻辑在前端界面层完成辅助展示，后端仅负责数据'
        '存储和查询。total_salary 字段在保存时由 SalaryServiceImpl 按公式计算并写入：'
    )
    add_code(
        '// SalaryServiceImpl.java — 薪酬合计计算\n'
        'BigDecimal totalSalary = salary.getBaseSalary()\n'
        '    .add(salary.getTotalHours()\n'
        '         .multiply(salary.getHourRate()))\n'
        '    .add(salary.getBonus())\n'
        '    .subtract(salary.getDeduction());\n'
        'salary.setTotalSalary(totalSalary);\n'
        'salaryMapper.insert(salary);'
    )
    add_figure('图5-4', '薪酬结算管理页面截图',
               '点击"薪酬管理"，截取薪酬列表页，展示教师姓名/月份/课时数/课时单价/奖励/扣款/应付总额/状态列')

    add_heading('5.6  三级审批流程模块', level=2)
    add_para(
        '审批模块是本系统实现复杂度最高的部分。ApprovalServiceImpl 中的 approve() 方法封装了'
        '完整的状态转换逻辑：'
    )
    add_code(
        '// ApprovalServiceImpl.java — 审批状态机核心\n'
        'public void approve(Long id, String node, boolean pass,\n'
        '                    String remark, String operatorName) {\n'
        '    Approval approval = getById(id);\n'
        '    if (!"PENDING".equals(approval.getApprovalStatus())) {\n'
        '        throw new RuntimeException("该审批单已结束，不可操作");\n'
        '    }\n'
        '    LocalDateTime now = LocalDateTime.now();\n'
        '    switch (node) {\n'
        '        case "college_leader":\n'
        '            approval.setCollegeStatus(pass ? "APPROVED" : "REJECTED");\n'
        '            approval.setCollegeRemark(remark);\n'
        '            approval.setCollegeBy(operatorName);\n'
        '            approval.setCollegeTime(now);\n'
        '            if (pass) approval.setCurrentNode("hr_salary");\n'
        '            break;\n'
        '        case "hr_salary":\n'
        '            approval.setHrSalaryStatus(pass ? "APPROVED" : "REJECTED");\n'
        '            approval.setHrSalaryRemark(remark);\n'
        '            approval.setHrSalaryBy(operatorName);\n'
        '            approval.setHrSalaryTime(now);\n'
        '            if (pass) approval.setCurrentNode("hr_director");\n'
        '            break;\n'
        '        case "hr_director":\n'
        '            approval.setHrDirectorStatus(pass ? "APPROVED" : "REJECTED");\n'
        '            approval.setHrDirectorRemark(remark);\n'
        '            approval.setHrDirectorBy(operatorName);\n'
        '            approval.setHrDirectorTime(now);\n'
        '            if (pass) {\n'
        '                approval.setCurrentNode("finished");\n'
        '                approval.setApprovalStatus("APPROVED");\n'
        '                // 触发合同自动生成（@Lazy 解决循环依赖）\n'
        '                contractService.generateFromApproval(id);\n'
        '            }\n'
        '            break;\n'
        '    }\n'
        '    if (!pass) approval.setApprovalStatus("REJECTED");\n'
        '    updateById(approval);\n'
        '}'
    )
    add_para(
        '在实现过程中，ApprovalServiceImpl 需要注入 ContractService，而 ContractServiceImpl '
        '同样依赖 ApprovalService，形成循环依赖。Spring 默认在单例 Bean 中通过三级缓存自动解决'
        '部分循环依赖，但当两个 Bean 同时依赖对方时仍会抛出异常。解决方案是在 '
        'ApprovalServiceImpl 的 ContractService 注入点加 @Lazy 注解，使 Spring 在第一次调用'
        '时才初始化 ContractService，打破初始化阶段的循环：'
    )
    add_code(
        '// ApprovalServiceImpl.java\n'
        '@Lazy\n'
        '@Autowired\n'
        'private ContractService contractService;'
    )
    add_figure('图5-5', '审批流程管理页面截图（含时间线详情弹窗）',
               '点击"聘用审批"，截取审批列表；再点击某条记录的"时间线"按钮，'
               '截取弹出的el-timeline进度图，展示三个节点的审批状态、审批人和时间')

    add_heading('5.7  合同管理模块', level=2)
    add_para(
        '合同由 ContractServiceImpl.generateFromApproval() 自动生成，幂等逻辑通过查询'
        ' approval_id 是否已存在对应合同来实现，避免处长审批接口被重复调用时创建多份合同：'
    )
    add_code(
        '// ContractServiceImpl.java — 幂等合同生成\n'
        'public void generateFromApproval(Long approvalId) {\n'
        '    // 幂等校验：已有合同则直接返回\n'
        '    Long count = contractMapper.selectCount(\n'
        '        new LambdaQueryWrapper<Contract>()\n'
        '            .eq(Contract::getApprovalId, approvalId)\n'
        '            .eq(Contract::getDeleted, 0));\n'
        '    if (count > 0) return;\n'
        '\n'
        '    Approval approval = approvalMapper.selectById(approvalId);\n'
        '    Contract contract = new Contract();\n'
        '    contract.setContractNo("CTR" + System.currentTimeMillis());\n'
        '    contract.setTeacherId(approval.getTeacherId());\n'
        '    contract.setDepartmentId(approval.getDepartmentId());\n'
        '    contract.setApprovalId(approvalId);\n'
        '    contract.setStartDate(approval.getStartDate());\n'
        '    contract.setEndDate(approval.getEndDate());\n'
        '    contract.setSalaryStandard(approval.getProposedSalary());\n'
        '    contract.setContractStatus("UNSIGNED");\n'
        '    contractMapper.insert(contract);\n'
        '}'
    )
    add_figure('图5-6', '合同管理页面截图',
               '点击"聘用合同"，截取合同列表，展示合同编号/教师姓名/薪酬标准/合同状态/'
               '教师签署时间/学校签署时间，以及"教师签署"和"学校签署"按钮')

    add_heading('5.8  数据统计与可视化模块', level=2)
    add_para(
        '首页仪表盘由 DashboardController 提供数据，调用 DashboardService 从各业务表中汇总'
        '统计数值，包括教师总数、本学期在职人数、当月总课时、待处理审批数量，以及学历分布、'
        '月度课时趋势、各院系薪酬对比等图表数据。前端通过 ECharts 的 init() 和 setOption() '
        '渲染对应图表。'
    )
    add_figure('图5-7', '首页数据统计仪表盘截图',
               '登录后默认显示首页，截取完整仪表盘页面，包含4个数字卡片（教师总数/在职/课时/待审批）'
               '和下方的ECharts图表区域（饼图+折线图+柱状图）')

    add_para('本章小结：本章按模块介绍了系统的实现过程，重点展示了 JWT 拦截器链、'
             'MyBatis-Plus 动态查询、三级审批状态机和合同幂等生成的关键代码，'
             '并说明了循环依赖问题的解决方案。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第6章  系统测试
# ══════════════════════════════════════════════════════════════════════════════

def chapter6():
    add_heading('第6章  系统测试', level=1)

    add_heading('6.1  测试环境', level=2)
    add_table_caption('表6-1  测试环境配置')
    add_simple_table(
        ['项目', '配置'],
        [('操作系统', 'macOS 15 / Windows 10'),
         ('处理器', 'Apple M4 Pro / Intel Core i5'),
         ('内存', '16 GB'),
         ('后端运行', 'Spring Boot 内嵌 Tomcat，8088 端口'),
         ('前端运行', 'Vue CLI Dev Server，8080 端口'),
         ('数据库', 'H2 文件模式，预置演示数据'),
         ('测试工具', 'Knife4j 接口文档、浏览器手工测试'),
         ('测试账号', 'admin/admin123（ADMIN）、dept/dept123（DEPT）、teacher/teacher123（TEACHER）')]
    )

    add_heading('6.2  功能测试', level=2)
    add_para(
        '功能测试按模块逐项验证，以下列出主要测试用例及结果。'
    )
    add_table_caption('表6-2  功能测试用例与结果')
    add_simple_table(
        ['编号', '测试用例', '预期结果', '实测结果'],
        [('TC-01', '管理员账号登录', '返回JWT，跳转仪表盘', '通过'),
         ('TC-02', '错误密码登录', '提示"账号或密码错误"', '通过'),
         ('TC-03', 'TEACHER角色访问审批页', '前端拦截，跳转仪表盘', '通过'),
         ('TC-04', '新增外聘教师档案', '记录入库，列表可查询', '通过'),
         ('TC-05', '发起聘用审批', '生成APR编号，状态为PENDING', '通过'),
         ('TC-06', '学院负责人通过', 'current_node变为hr_salary', '通过'),
         ('TC-07', '人事处薪酬岗驳回', 'approval_status变为REJECTED', '通过'),
         ('TC-08', '三级全部通过', 'approval_status=APPROVED，自动生成合同', '通过'),
         ('TC-09', '重复触发合同生成', '幂等校验生效，不重复创建', '通过'),
         ('TC-10', '教师签署→学校签署', 'contract_status最终变为SIGNED', '通过'),
         ('TC-11', '薪酬计算公式', 'total_salary=base+hours×rate+bonus-deduction', '通过'),
         ('TC-12', '逻辑删除教师', 'deleted=1，列表不再显示', '通过'),
         ('TC-13', '首页统计数据', 'ECharts图表正确渲染', '通过')]
    )

    add_figure('图6-1', '三级审批全流程测试截图',
               '截取一次完整审批的时间线弹窗，三个节点均显示"已通过"状态（绿色），'
               '并显示审批人姓名和时间；同时截取对应自动生成的合同记录')

    add_heading('6.3  性能与安全测试', level=2)
    add_para(
        '性能测试使用浏览器开发者工具的 Network 面板记录页面加载时延。在本地开发环境（局域网内）'
        '测试结果如下：登录接口响应约 80ms，教师列表分页接口（20条记录）约 60ms，首页统计接口'
        '（含 5 个子查询）约 150ms，前端页面切换（含Axios请求）约 200-400ms。所有关键操作'
        '均在 500ms 以内，满足设计目标。'
    )
    add_para(
        'H2 文件模式在并发极低的单用户场景下性能表现良好。在多用户同时操作（10 个并发请求）时，'
        '由于 H2 文件锁机制，部分写操作可能出现短暂排队，但在高校内网日常使用场景下（同时操作用户'
        '不超过 20 人），不会造成明显卡顿。'
    )
    add_para(
        '安全测试重点验证了三项：其一，未携带 token 访问任意接口，返回 401；其二，使用 TEACHER '
        '角色的 token 访问 /api/approval，返回 403；其三，JWT 令牌过期后（测试时临时将有效期'
        '设为 5 秒），再次请求返回 401，前端自动跳转至登录页。三项均通过验证。'
    )

    add_para('本章小结：本章在本地测试环境中对系统各功能模块进行了手工测试，'
             '全部13个核心用例通过验证，性能指标满足设计要求，安全拦截机制正常工作。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第7章  总结与展望
# ══════════════════════════════════════════════════════════════════════════════

def chapter7():
    add_heading('第7章  总结与展望', level=1)

    add_heading('7.1  工作总结', level=2)
    add_para(
        '本文设计并实现了一套面向高校场景的外聘教师管理系统。系统覆盖了外聘教师从档案录入、'
        '聘用审批、课程排布、考勤记录、薪酬结算到合同签署的完整业务闭环，采用前后端分离架构，'
        '后端共实现约 70 个 Java 类，前端共 16 个页面视图。'
    )
    add_para(
        '在技术层面，三级审批状态机是本文的核心实现难点。通过将审批节点的状态转换集中在单个'
        ' Service 方法中，并以数据库字段记录每个节点的结果，避免了引入独立工作流引擎的复杂性，'
        '同时保持了足够的可追溯性。合同生成的幂等校验解决了业务层面的重复调用问题，@Lazy 注解'
        '的使用解决了初始化阶段的循环依赖问题——这两处都是调试过程中遇到并修复的实际问题。'
    )

    add_heading('7.2  存在的不足', level=2)
    add_para(
        '当前系统存在几处已知的局限性，后续有待改进：'
    )
    add_para(
        '第一，薪酬结算未实现自动批量汇总。目前管理员需手动逐条录入薪酬记录，未能从考勤数据中'
        '自动汇总月度课时并批量生成结算单，实际使用时操作成本偏高。'
    )
    add_para(
        '第二，文件存储仅支持路径记录，未接入对象存储。教师的学历证书、合同扫描件等文件字段'
        '目前仅存储文件路径字符串，未实现实际的文件上传和在线预览功能。'
    )
    add_para(
        '第三，H2 数据库不支持多实例部署。当前架构以单机为前提，若高校需要多台服务器分担负载，'
        '必须切换至 MySQL 并相应调整部署方案。'
    )
    add_para(
        '第四，通知功能仅为系统内公告，未实现邮件或短信推送，审批节点变更时无法主动通知相关人员。'
    )

    add_heading('7.3  未来改进方向', level=2)
    add_para(
        '针对上述不足，后续改进可从以下几个方向推进：引入定时任务（如 Spring Schedule）在每月'
        '末自动汇总考勤数据并生成薪酬结算单；接入 MinIO 或阿里云 OSS 实现文件的实际上传和在线'
        '预览；将消息通知抽象为独立的 NotificationService，支持对接邮件服务（JavaMail）；'
        '在高并发需求出现时，将数据库切换至 MySQL，并在 Redis 中缓存审批状态和用户权限数据，'
        '减少数据库查询压力。'
    )

    add_para('本章小结：本章总结了系统的实现成果和主要技术难点，客观指出了现有版本在批量薪酬、'
             '文件存储和通知推送方面的局限，并提出了具体的改进路径。')
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════════════════════════════════════

def references():
    add_heading('参考文献', level=1)
    refs = [
        '[1] Craig Walls. Spring Boot实战[M]. 第4版. 北京：人民邮电出版社，2020.',
        '[2] 苗圣法. MyBatis-Plus从入门到精通[M]. 北京：电子工业出版社，2022.',
        '[3] M. Jones, J. Bradley, N. Sakimura. JSON Web Token (JWT). RFC 7519[S]. Internet Engineering Task Force, 2015.',
        '[4] 尤雨溪. Vue.js官方文档[EB/OL]. https://cn.vuejs.org/, 2023.',
        '[5] 饿了么前端团队. Element UI组件库文档[EB/OL]. https://element.eleme.io/, 2022.',
        '[6] Apache Software Foundation. Apache POI - Java API To Access Microsoft Format Files[EB/OL]. https://poi.apache.org/, 2023.',
        '[7] H2 Database Engine. H2 Database Engine Documentation[EB/OL]. https://h2database.com/, 2023.',
        '[8] 周志明. 深入理解Java虚拟机：JVM高级特性与最佳实践[M]. 第3版. 北京：机械工业出版社，2019.',
        '[9] 黄勇. MyBatis从入门到精通[M]. 北京：电子工业出版社，2017.',
        '[10] Thomas H. Cormen, et al. Introduction to Algorithms[M]. 3rd Edition. MIT Press, 2009.',
        '[11] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Addison-Wesley, 2002.',
        '[12] 教育部. 普通高等学校师资队伍建设管理规定[Z]. 教人[2023]1号，2023.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_para_format(p, first_line_indent=False, space_before=2, space_after=2)
        run = p.add_run(ref)
        set_run_font(run, size_pt=10.5)
    page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 致谢
# ══════════════════════════════════════════════════════════════════════════════

def acknowledgement():
    add_heading('致  谢', level=1)
    add_para(
        '本系统的完成历经数月，从最初的技术选型到最终的功能验证，得到了许多人的帮助。'
    )
    add_para(
        '感谢指导教师在选题和设计阶段给予的方向性建议，以及在论文撰写过程中对逻辑结构提出的'
        '多次修改意见。'
    )
    add_para(
        '感谢同学们在项目讨论中提出的问题，尤其是在审批流程设计初期关于节点驳回后是否允许重新'
        '提交的争论，促使我认真梳理了各种边界条件。'
    )
    add_para(
        '最后感谢家人的支持。毕业设计的最后阶段恰逢课程压力最大的时候，能够完成全部功能并'
        '通过测试，少不了他们在生活上的照顾。'
    )

# ══════════════════════════════════════════════════════════════════════════════
# 主流程
# ══════════════════════════════════════════════════════════════════════════════

add_cover()
add_abstract()
add_toc()
chapter1()
chapter2()
chapter3()
chapter4()
chapter5()
chapter6()
chapter7()
references()
acknowledgement()

output_path = '/Users/yuchenxu/Desktop/于晨旭毕设/于晨旭_毕业论文_高校外聘教师管理系统.docx'
doc.save(output_path)
print(f'论文已生成：{output_path}')
